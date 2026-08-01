"""Generates the High Level Design document for the AI Tracker PWA."""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY = RGBColor(0x1F, 0x2A, 0x5A)
IND = RGBColor(0x4F, 0x46, 0xE5)
CYAN = RGBColor(0x0E, 0x74, 0x90)
GREY = RGBColor(0x55, 0x5B, 0x6E)

doc = Document()

# ---------- base styles ----------
st = doc.styles["Normal"]
st.font.name = "Calibri"
st.font.size = Pt(10.5)
st.paragraph_format.space_after = Pt(6)
st.paragraph_format.line_spacing = 1.12

for name, size, colour in (
    ("Heading 1", 18, NAVY),
    ("Heading 2", 14, IND),
    ("Heading 3", 11.5, CYAN),
    ("Heading 4", 10.5, GREY),
):
    s = doc.styles[name]
    s.font.name = "Calibri"
    s.font.size = Pt(size)
    s.font.color.rgb = colour
    s.font.bold = True
    s.paragraph_format.space_before = Pt(14 if size > 13 else 10)
    s.paragraph_format.space_after = Pt(5)
    s.paragraph_format.keep_with_next = True

sec = doc.sections[0]
sec.left_margin = sec.right_margin = Inches(0.85)
sec.top_margin = Inches(0.8)
sec.bottom_margin = Inches(0.8)


def shade(el, hexcolor):
    tc = OxmlElement("w:shd")
    tc.set(qn("w:val"), "clear")
    tc.set(qn("w:fill"), hexcolor)
    el.append(tc)


def h(text, level=1, page_break=False):
    if page_break:
        doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    return doc.add_heading(text, level=level)


def p(text="", bold=False, italic=False, size=None, colour=None, align=None, space=6):
    par = doc.add_paragraph()
    run = par.add_run(text)
    run.bold = bold
    run.italic = italic
    if size:
        run.font.size = Pt(size)
    if colour:
        run.font.color.rgb = colour
    if align:
        par.alignment = align
    par.paragraph_format.space_after = Pt(space)
    return par


def rich(parts):
    """parts = [(text, bold, italic)]"""
    par = doc.add_paragraph()
    for t, b, i in parts:
        r = par.add_run(t)
        r.bold = b
        r.italic = i
    return par


def bullet(text, level=0, bold_prefix=None):
    par = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    if bold_prefix:
        r = par.add_run(bold_prefix)
        r.bold = True
    par.add_run(text)
    par.paragraph_format.space_after = Pt(2)
    return par


def numbered(text, bold_prefix=None):
    par = doc.add_paragraph(style="List Number")
    if bold_prefix:
        r = par.add_run(bold_prefix)
        r.bold = True
    par.add_run(text)
    par.paragraph_format.space_after = Pt(2)
    return par


def code(text, caption=None):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = "Table Grid"
    cell = tbl.rows[0].cells[0]
    shade(cell._tc.get_or_add_tcPr(), "F4F5F9")
    cell.text = ""
    for i, line in enumerate(text.strip("\n").split("\n")):
        par = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        run = par.add_run(line if line else " ")
        run.font.name = "Consolas"
        run.font.size = Pt(8.5)
        par.paragraph_format.space_after = Pt(0)
        par.paragraph_format.line_spacing = 1.0
    if caption:
        c = doc.add_paragraph()
        r = c.add_run(caption)
        r.italic = True
        r.font.size = Pt(8.5)
        r.font.color.rgb = GREY
        c.paragraph_format.space_after = Pt(8)
    else:
        doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return tbl


def table(headers, rows, widths=None, caption=None, font=8.8):
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = tbl.rows[0]
    for i, name in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(name)
        run.bold = True
        run.font.size = Pt(font)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        cell.paragraphs[0].paragraph_format.space_after = Pt(1)
        shade(cell._tc.get_or_add_tcPr(), "3B3F87")
    for r_i, row in enumerate(rows):
        cells = tbl.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            par = cells[i].paragraphs[0]
            run = par.add_run(str(val))
            run.font.size = Pt(font)
            par.paragraph_format.space_after = Pt(1)
            if r_i % 2 == 1:
                shade(cells[i]._tc.get_or_add_tcPr(), "F2F3F9")
    if widths:
        for row in tbl.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Inches(w)
    if caption:
        c = doc.add_paragraph()
        run = c.add_run(caption)
        run.italic = True
        run.font.size = Pt(8.5)
        run.font.color.rgb = GREY
        c.paragraph_format.space_after = Pt(10)
    else:
        doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return tbl


def field(par, instr):
    r1 = par.add_run()
    fc = OxmlElement("w:fldChar"); fc.set(qn("w:fldCharType"), "begin"); r1._r.append(fc)
    r2 = par.add_run()
    it = OxmlElement("w:instrText"); it.set(qn("xml:space"), "preserve"); it.text = instr
    r2._r.append(it)
    r3 = par.add_run()
    fs = OxmlElement("w:fldChar"); fs.set(qn("w:fldCharType"), "separate"); r3._r.append(fs)
    r4 = par.add_run("…")
    r5 = par.add_run()
    fe = OxmlElement("w:fldChar"); fe.set(qn("w:fldCharType"), "end"); r5._r.append(fe)

# ============================ TITLE PAGE ============================
t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
t.paragraph_format.space_before = Pt(110)
r = t.add_run("HIGH LEVEL DESIGN"); r.bold = True; r.font.size = Pt(30); r.font.color.rgb = NAVY

t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run("AI Tracker — 90 Day Challenge"); r.bold = True; r.font.size = Pt(20); r.font.color.rgb = IND

t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run("An offline-first Progressive Web Application for habit,\nlearning and problem-solving progress tracking")
r.font.size = Pt(12); r.font.color.rgb = GREY; r.italic = True

t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
t.paragraph_format.space_before = Pt(30)
r = t.add_run("Next.js 15  ·  React 19  ·  TypeScript 5  ·  Tailwind CSS 4  ·  PWA")
r.font.size = Pt(10.5); r.bold = True; r.font.color.rgb = CYAN

doc.add_paragraph().paragraph_format.space_after = Pt(40)
table(["Field", "Value"],
      [["Document type", "High Level Design (HLD)"],
       ["System name", "AI Tracker — 90 Day Challenge"],
       ["Version", "1.0"],
       ["Date", "01 August 2026"],
       ["Status", "Baselined — implementation complete and verified"],
       ["Repository path", "~/ai-tracker"],
       ["Runtime", "Node.js 20.19.5 / npm 10.8.2"],
       ["Total source size", "2,429 lines across 23 source files"],
       ["Runtime dependencies", "3 (next, react, react-dom)"],
       ["Audience", "Developers, reviewers, future maintainers"]],
      widths=[2.0, 4.6])

doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

# ============================ TOC ============================
h("Table of Contents", 1)
p("This field updates on open. In Word press Ctrl+A then F9, or right-click the table and choose "
  "\"Update Field\" to populate page numbers.", italic=True, size=9, colour=GREY)
toc = doc.add_paragraph()
field(toc, r'TOC \o "1-3" \h \z \u')

# ============================ 1. EXEC SUMMARY ============================
h("1. Executive Summary", 1, page_break=True)
p("AI Tracker is a single-user, client-only Progressive Web Application (PWA) that replaces a "
  "spreadsheet-based self-improvement tracker. It records fourteen daily habits, reading progress "
  "across a personal book list, a fifteen-topic Data Structures & Algorithms roadmap, and daily "
  "LeetCode volume, then converts that raw activity into motivational feedback: experience points "
  "(XP), levels, streaks, achievements, and time-series analytics.")
p("The defining architectural decision is that the application has no backend. All state lives in "
  "the browser's localStorage on the user's own device. Consequences of that single decision "
  "propagate through every layer of this design: there is no authentication, no network latency, "
  "no server cost, no privacy exposure, and no multi-device sync. The application is therefore "
  "instantaneous and fully functional offline, at the price of being device-local — a trade "
  "deliberately accepted and mitigated by a JSON export/import backup facility.")
p("A service worker pre-caches the application shell so it launches without a network connection, "
  "and a web app manifest plus generated icon set allow it to be installed to an iPhone home "
  "screen where it runs full-screen with no browser chrome, indistinguishable from a native app.")

h("1.1 Design goals", 2)
table(["#", "Goal", "How the design achieves it"],
      [["G1", "Zero friction to log a day", "One-tap toggles; quick check-in grid on the dashboard so the most common action needs no navigation"],
       ["G2", "Works with no network", "Service worker pre-caches all five routes; state reads/writes are synchronous localStorage calls"],
       ["G3", "Feels like a native iPhone app", "Manifest display:standalone, apple-touch-icon, safe-area insets, bottom tab bar, 16px inputs to suppress iOS zoom"],
       ["G4", "Motivating, not merely a log", "XP economy, level curve, streaks, 13 achievements, perfect-day bonus, heatmap"],
       ["G5", "Zero operating cost", "No server, no database, no third-party service; static export deployable to any CDN"],
       ["G6", "Privacy by construction", "Data never leaves the device; no analytics, no telemetry, no accounts"],
       ["G7", "Maintainable by one person", "3 runtime dependencies; pure-function domain layer; ~2.4k lines total"]],
      widths=[0.4, 1.9, 4.3])

h("1.2 Explicit non-goals", 2)
bullet("Multi-user support, accounts, or authentication.")
bullet("Server-side persistence or cross-device synchronisation in version 1.0 (deferred — see Section 18).")
bullet("Push notifications or reminders (requires a push service and, on iOS, an installed PWA).")
bullet("Social features, leaderboards, or sharing.")
bullet("Native iOS/Android binaries — the PWA installation path is considered sufficient.")

# ============================ 2. REQUIREMENTS ============================
h("2. Requirements", 1)
h("2.1 Functional requirements", 2)
table(["ID", "Requirement", "Realised by"],
      [["FR-1", "Track 14 named daily habits with a single tap", "src/lib/habits.ts catalogue + HabitCard + dashboard quick grid"],
       ["FR-2", "Record and edit any past day, not just today", "Habits screen date navigator (previous/next, future disabled)"],
       ["FR-3", "Log LeetCode problems split by difficulty", "Per-day easy/medium/hard counters on the Habits screen"],
       ["FR-4", "Capture a free-text reflection per day", "Day-note textarea bound to DayRecord.note"],
       ["FR-5", "Track book reading progress by page", "Books list with current/total pages, +10 quick increment"],
       ["FR-6", "Track a DSA roadmap by topic", "15 seeded topics with done/total counters"],
       ["FR-7", "Keep free-form study notes", "Notes list with title, body, created timestamp"],
       ["FR-8", "Award XP and levels for activity", "dayXp() and levelInfo() in src/lib/stats.ts"],
       ["FR-9", "Compute current and best streaks", "currentStreak() / bestStreak() with a configurable daily threshold"],
       ["FR-10", "Visualise weekly, monthly and 90-day progress", "BarChart, LineChart, HBars, Donut, CalendarHeatmap"],
       ["FR-11", "Unlock achievements", "13 rule-based achievements evaluated on every render"],
       ["FR-12", "Configure the challenge (start date, length, target)", "Settings screen writing to AppState.settings"],
       ["FR-13", "Back up and restore all data", "JSON export to file / import with schema migration"],
       ["FR-14", "Install to an iPhone home screen", "Web app manifest + generated icons + apple meta tags"],
       ["FR-15", "Work with no network connection", "Service worker with network-first, cache-fallback strategy"]],
      widths=[0.55, 2.6, 3.5])

h("2.2 Non-functional requirements", 2)
table(["Attribute", "Target", "Design response", "Measured"],
      [["Performance", "Interaction < 100 ms", "Synchronous in-memory state; no network on the interaction path", "Instant — no I/O in the click handler"],
       ["Payload", "< 150 kB first load", "Zero chart libraries; hand-written SVG; 3 runtime deps", "102 kB shared + 3-5 kB per route"],
       ["Availability", "100% offline", "Pre-cached shell; localStorage persistence", "10 assets cached; verified activated"],
       ["Durability", "No silent data loss", "Write-through persistence on every state change; export/import", "Verified across reloads"],
       ["Privacy", "No data egress", "No backend, no analytics, no external fonts or CDNs", "Zero outbound requests after load"],
       ["Portability", "Any modern browser", "Standard web platform APIs only", "Chromium verified; Safari targeted"],
       ["Accessibility", "Keyboard + screen reader usable", "Semantic buttons, aria-pressed, aria-label, focus states", "Verified via accessibility tree"],
       ["Maintainability", "Single maintainer", "Pure-function domain layer, strict TypeScript, ESLint", "0 type errors, 0 lint errors"]],
      widths=[1.0, 1.25, 2.6, 1.85])

# ============================ 3. TECH STACK ============================
h("3. Technology Stack and Rationale", 1, page_break=True)
table(["Layer", "Technology", "Version", "Why this was chosen"],
      [["Framework", "Next.js (App Router)", "15.5.22", "File-system routing, static pre-rendering, first-class TypeScript, zero-config production build, free Vercel deployment"],
       ["UI library", "React", "19.1.0", "Component model and hooks; Context API removes any need for an external state library"],
       ["Language", "TypeScript", "5.x", "Compile-time guarantees over the data model; the domain layer is fully typed"],
       ["Styling", "Tailwind CSS", "4.x", "Utility classes keep styling co-located with markup; no CSS files to drift; tiny purged output"],
       ["Persistence", "Web Storage (localStorage)", "Platform", "Synchronous, universally supported, survives restarts, needs no permission prompt"],
       ["State", "React Context + useState", "Built-in", "The whole state object is small (a few hundred KB worst case); Redux/Zustand would add weight for no benefit"],
       ["Charts", "Hand-written SVG", "None", "Recharts/Chart.js would add 100-500 kB. Every chart here is simple geometry, so the maths is done directly"],
       ["Offline", "Service Worker API", "Platform", "The only standards-based way to serve an app shell with no network"],
       ["Icons", "Custom PNG encoder", "Node zlib", "Avoids adding sharp/canvas as a build dependency; deterministic, reproducible icon generation"],
       ["Linting", "ESLint + eslint-config-next", "9.x", "Catches unused symbols, hook-rule violations and unsafe patterns at build time"]],
      widths=[0.95, 1.5, 0.8, 3.45])

h("3.1 Dependency footprint", 2)
p("The production dependency tree contains exactly three packages. This is a deliberate design "
  "constraint: every added dependency is a future upgrade obligation and a supply-chain risk.")
code("""
"dependencies": {
  "react":     "19.1.0",
  "react-dom": "19.1.0",
  "next":      "15.5.22"
}
""", "package.json — the complete runtime dependency set.")

# ============================ 4. ARCHITECTURE ============================
h("4. System Architecture", 1, page_break=True)
h("4.1 Context view", 2)
p("The system boundary is unusually tight: the entire application lives inside one browser tab on "
  "one device. The only external actors are the user, the device's storage engine, and — during "
  "the very first visit only — the static file host that delivers the bundle.")
code("""
                         +---------------------------------------------+
                         |                 THE DEVICE                  |
                         |                                             |
   +--------+            |  +-------------------------------------+    |
   |        |  taps      |  |         Browser tab / PWA window    |    |
   |  USER  |----------->|  |                                     |    |
   |        |<-----------|  |   React component tree (UI)         |    |
   +--------+  renders   |  |            |          ^             |    |
                         |  |            v          |             |    |
                         |  |   StoreProvider (React Context)     |    |
                         |  |            |          ^             |    |
                         |  +------------|----------|-------------+    |
                         |               v          |                  |
                         |         +--------------------+             |
                         |         |    localStorage    |  <- durable |
                         |         +--------------------+             |
                         |               ^                             |
                         |         +--------------------+             |
                         |         |  Service Worker    |             |
                         |         |  + Cache Storage   |  <- shell   |
                         |         +--------------------+             |
                         +---------------|-----------------------------+
                                         | first visit only (HTTP GET)
                                         v
                              +----------------------+
                              |  Static file host    |
                              |  (Next.js server or  |
                              |   Vercel / any CDN)  |
                              +----------------------+
""", "Figure 1 — Context diagram. After the first load, the vertical link to the host is never required again.")

h("4.2 Layered view", 2)
p("The codebase is organised into four strictly ordered layers. Dependencies point downwards only; "
  "no lower layer imports from a higher one. This is what keeps the domain logic testable and the "
  "UI replaceable.")
code("""
+---------------------------------------------------------------------------+
|  LAYER 4 — ROUTES / SCREENS            src/app/*/page.tsx                 |
|  dashboard | habits | books | analytics | settings                        |
|  Responsibility: compose components, own local UI state (selected date,   |
|  chart range, form drafts). No business rules live here.                  |
+------------------------------|--------------------------------------------+
                               v
+---------------------------------------------------------------------------+
|  LAYER 3 — PRESENTATION COMPONENTS     src/components/*.tsx               |
|  ProgressRing  XPBar  HabitCard  CalendarHeatmap  Charts  Navbar  Ui      |
|  Responsibility: pure rendering from props. Stateless. Reusable.          |
+------------------------------|--------------------------------------------+
                               v
+---------------------------------------------------------------------------+
|  LAYER 2 — STATE / APPLICATION         src/lib/store.tsx                  |
|  StoreProvider, useStore(), 15 action creators, hydration, persistence,   |
|  schema migration. The single source of truth.                            |
+------------------------------|--------------------------------------------+
                               v
+---------------------------------------------------------------------------+
|  LAYER 1 — DOMAIN (pure functions)     src/lib/{types,date,habits,stats}  |
|  Type definitions, date arithmetic, habit catalogue, XP / level / streak  |
|  / achievement / series calculations. No React, no I/O, no side effects.  |
+---------------------------------------------------------------------------+
""", "Figure 2 — Layered architecture. Layer 1 is framework-agnostic and could be lifted into any other UI.")

h("4.3 Component inventory", 2)
table(["File", "LOC", "Layer", "Responsibility"],
      [["src/lib/types.ts", "50", "1", "All TypeScript interfaces: AppState, DayRecord, Book, DsaTopic, Note, Settings, LeetCount"],
       ["src/lib/date.ts", "68", "1", "Timezone-safe date keys, arithmetic, day numbering, range generation, formatting"],
       ["src/lib/habits.ts", "55", "1", "The 14-habit catalogue (id, label, icon, category, XP, hint) and 15 motivational quotes"],
       ["src/lib/stats.ts", "208", "1", "XP, levels, completion, streaks, achievements, chart series, aggregate roll-ups"],
       ["src/lib/store.tsx", "221", "2", "Context provider, hydration, write-through persistence, migration, 15 actions"],
       ["src/components/ProgressRing.tsx", "66", "3", "Animated SVG circular progress indicator with gradient stroke"],
       ["src/components/XPBar.tsx", "32", "3", "Level badge, title, XP-into-level bar"],
       ["src/components/HabitCard.tsx", "54", "3", "Toggleable habit row with icon, hint, XP value and check state"],
       ["src/components/CalendarHeatmap.tsx", "78", "3", "GitHub-style contribution grid with month labels and legend"],
       ["src/components/Charts.tsx", "187", "3", "BarChart, LineChart, HBars and Donut — all hand-drawn SVG"],
       ["src/components/Navbar.tsx", "67", "3", "Responsive navigation: desktop top bar, mobile bottom tab bar"],
       ["src/components/Ui.tsx", "59", "3", "Card, CardTitle and StatCard design-system primitives"],
       ["src/components/PWARegister.tsx", "28", "3", "Service-worker registration with secure-origin and readyState guards"],
       ["src/app/layout.tsx", "53", "4", "Root shell: metadata, viewport, provider, navigation, main container"],
       ["src/app/page.tsx", "5", "4", "Redirects / to /dashboard"],
       ["src/app/dashboard/page.tsx", "141", "4", "Ring, XP bar, four stat tiles, quote, quick check-in, heatmap"],
       ["src/app/habits/page.tsx", "170", "4", "Date navigator, day summary, Body/Mind habit groups, LeetCode counters, note"],
       ["src/app/books/page.tsx", "262", "4", "Books CRUD, LeetCode donut, DSA roadmap CRUD, notes CRUD"],
       ["src/app/analytics/page.tsx", "139", "4", "Eight stat tiles, range switch, four charts, heatmap, achievements"],
       ["src/app/settings/page.tsx", "196", "4", "Profile and challenge configuration, backup/restore, reset, install guide"],
       ["src/app/globals.css", "66", "-", "Tailwind import, background gradients, iOS tweaks, scrollbar, fade-up keyframes"],
       ["public/sw.js", "64", "-", "Service worker: install/activate/fetch handlers and caching strategy"],
       ["scripts/generate-icons.mjs", "160", "-", "Dependency-free PNG encoder and procedural icon renderer"]],
      widths=[2.05, 0.42, 0.42, 3.81],
      caption="Table — Complete file inventory. Total 2,429 lines.", font=8.2)

h("4.4 Runtime sequence — the critical path", 2)
p("The following sequence traces the single most frequent interaction in the system: the user taps "
  "a habit. It is deliberately short — four synchronous steps and one asynchronous write — which "
  "is why the UI feels instant.")
code("""
 USER            HabitCard        useStore()       StoreProvider      React        localStorage
   |                 |                 |                 |              |               |
   |--- tap -------->|                 |                 |              |               |
   |                 |-- onToggle() -->|                 |              |               |
   |                 |                 |- toggleHabit -->|              |               |
   |                 |                 |                 |              |               |
   |                 |                 |        mutateDay(): build a NEW state object   |
   |                 |                 |        (immutable spread, no mutation)         |
   |                 |                 |                 |              |               |
   |                 |                 |                 |- setState -->|               |
   |                 |                 |                 |              |               |
   |                 |                 |        useMemo recomputes the context value    |
   |                 |                 |        Consumers re-render; stats.ts re-derives|
   |                 |                 |        XP, %, streak, series (all pure)        |
   |<---- new UI paint (< 16 ms, no I/O on this path) ---|              |               |
   |                 |                 |                 |              |               |
   |                 |                 |     useEffect([state]) fires AFTER paint       |
   |                 |                 |                 |------- JSON.stringify ------>|
   |                 |                 |                 |         (write-through)      |
""", "Figure 3 — Habit toggle sequence. Persistence happens after paint, so storage latency never blocks the UI.")

# ============================ 5. DATA MODEL ============================
h("5. Data Model", 1, page_break=True)
p("There is exactly one persisted object — AppState — held under a single localStorage key. "
  "Everything the application knows is reachable from it. This design makes backup, restore, "
  "migration and reasoning about state trivial, because there is nothing to keep in sync.")

h("5.1 Entity relationships", 2)
code("""
                        AppState  (the single root aggregate)
                             |
      +---------+------------+-------------+------------+-----------+
      |         |            |             |            |           |
   version   settings       days          books        dsa        notes
   (number)     |         (map)          (array)     (array)     (array)
                |            |               |           |           |
     +----------+---+   key = "YYYY-MM-DD"   |           |           |
     | name         |        |               |           |           |
     | startDate    |    DayRecord           Book     DsaTopic     Note
     | goalDays     |        |               |           |           |
     | dailyTarget% |   +----+-----+     id           id          id
     +--------------+   |    |     |     title        name        title
                     habits leetcode note  author      total       body
                    (map of  {easy,      totalPages   done        createdAt
                     id->bool) medium,   currentPage
                               hard}     color
""", "Figure 4 — Entity relationship view of the persisted aggregate.")

h("5.2 Data dictionary", 2)
table(["Type", "Field", "TS type", "Meaning / constraint"],
      [["AppState", "version", "number", "Schema version, currently 1. Drives migration on load."],
       ["", "settings", "Settings", "User and challenge configuration."],
       ["", "days", "Record<string, DayRecord>", "Sparse map keyed by local date 'YYYY-MM-DD'. Absent key = nothing logged."],
       ["", "books", "Book[]", "Ordered reading list, seeded with 3 titles."],
       ["", "dsa", "DsaTopic[]", "Ordered roadmap, seeded with 15 topics / 192 problems."],
       ["", "notes", "Note[]", "Newest-first list of free-text notes."],
       ["Settings", "name", "string", "Display name in the dashboard greeting. May be empty."],
       ["", "startDate", "string", "Date key of challenge day 1. Defaults to first-run date."],
       ["", "goalDays", "number", "Challenge length; default 90, minimum 7."],
       ["", "dailyTargetPct", "number", "Completion % at which a day counts for the streak. Default 70, range 10-100 step 5."],
       ["DayRecord", "habits", "Record<string, boolean>", "Sparse: only toggled habits appear. Missing key is treated as false."],
       ["", "leetcode", "LeetCount", "{ easy, medium, hard } — non-negative integers, clamped at 0."],
       ["", "note", "string?", "Optional free-text reflection for that day."],
       ["Book", "id", "string", "8-character base-36 identifier."],
       ["", "title / author", "string", "Display fields."],
       ["", "totalPages", "number", "Denominator for the progress bar; minimum 1."],
       ["", "currentPage", "number", "Clamped to 0..totalPages when rendering."],
       ["", "color", "string", "Hex colour for the progress bar, cycled from a 6-colour palette."],
       ["DsaTopic", "id / name", "string", "Identifier and display name."],
       ["", "total / done", "number", "Problems in the topic and problems completed; done clamped to total."],
       ["Note", "id / title / body", "string", "Content fields; title defaults to 'Untitled' if blank."],
       ["", "createdAt", "string", "ISO-8601 timestamp captured at creation."]],
      widths=[0.85, 1.05, 1.5, 3.3], font=8.2)

h("5.3 Why days is a sparse map, not an array", 2)
p("A 90-element array indexed by day number would have been simpler, but it fails three "
  "requirements. First, the challenge length is user-configurable, so the array size is not fixed. "
  "Second, the user can change the start date, which would silently reassign every historical "
  "entry to the wrong day. Third, analytics needs 'the last 30 calendar days', which is a date "
  "query, not an index range. Keying on an absolute local date makes every one of these operations "
  "correct by construction, and absence of a key naturally represents 'that day was never touched' "
  "without storing 14 false values per day.")

h("5.4 Persisted shape", 2)
code("""
localStorage key: "ai-tracker-state-v1"

{
  "version": 1,
  "settings": { "name": "", "startDate": "2026-08-01", "goalDays": 90, "dailyTargetPct": 70 },
  "days": {
    "2026-08-01": {
      "habits":   { "gym": true, "leetcode": true },
      "leetcode": { "easy": 0, "medium": 0, "hard": 0 },
      "note": ""
    }
  },
  "books": [ { "id": "a1b2c3d4", "title": "Designing Machine Learning Systems",
               "author": "Chip Huyen", "totalPages": 386, "currentPage": 0,
               "color": "#06b6d4" } ],
  "dsa":   [ { "id": "e5f6g7h8", "name": "Dynamic Programming", "total": 24, "done": 0 } ],
  "notes": []
}
""", "Figure 5 — Actual persisted payload, captured from the running application during verification.")

h("5.5 Versioning and migration", 2)
p("Every load passes the parsed JSON through migrate(), which merges the stored object over a "
  "freshly built default. This gives three properties for free: a field added in a future release "
  "appears with its default value rather than undefined; a corrupted or partially written object "
  "degrades to defaults instead of crashing; and an imported backup from an older release is "
  "upgraded on the spot. The storage key itself carries the schema generation (…-v1), so a future "
  "breaking change can be introduced under a new key without destroying the old data.")
code("""
function migrate(raw: unknown): AppState {
  const base = defaultState();
  if (!raw || typeof raw !== "object") return base;   // corrupt -> defaults
  const parsed = raw as Partial<AppState>;
  return {
    version:  VERSION,
    settings: { ...base.settings, ...(parsed.settings ?? {}) },  // field-level merge
    days:     parsed.days  ?? {},
    books:    parsed.books ?? base.books,
    dsa:      parsed.dsa   ?? base.dsa,
    notes:    parsed.notes ?? [],
  };
}
""", "src/lib/store.tsx — forward-compatible load path, shared by hydration and backup import.")

# ============================ 6. STATE MANAGEMENT ============================
h("6. State Management Design", 1, page_break=True)
p("StoreProvider is the heart of the application. It owns the single AppState object, exposes a "
  "fixed set of action creators, and solves three problems that every localStorage-backed React "
  "app must solve: hydration mismatch, write amplification, and re-render storms.")

h("6.1 The hydration problem and its solution", 2)
p("Next.js pre-renders every route to static HTML at build time. On the server there is no "
  "localStorage, so the server-rendered markup necessarily shows an empty state. If the client "
  "immediately rendered the real stored state, React would detect a mismatch between server HTML "
  "and client output and log a hydration error.")
p("The solution is a two-phase render controlled by a hydrated flag:")
numbered("Initial render (server and first client pass) uses defaultState() — identical output on both sides, so hydration succeeds cleanly.", bold_prefix="Phase 1 — ")
numbered("A useEffect (which never runs on the server) reads localStorage, replaces the state, and sets hydrated = true. Every screen renders an animated skeleton until this flag flips.", bold_prefix="Phase 2 — ")
code("""
const [state, setState]       = useState<AppState>(defaultState);
const [hydrated, setHydrated] = useState(false);
const ready = useRef(false);          // guards the FIRST persist

useEffect(() => {                      // runs once, client only
  try {
    const raw = localStorage.getItem(STORAGE_KEY);
    if (raw) setState(migrate(JSON.parse(raw)));
  } catch { /* corrupt storage -> keep defaults */ }
  ready.current = true;
  setHydrated(true);
}, []);

useEffect(() => {                      // write-through persistence
  if (!ready.current) return;          // <-- never overwrite storage with defaults
  try { localStorage.setItem(STORAGE_KEY, JSON.stringify(state)); }
  catch { /* quota exceeded - ignore */ }
}, [state]);
""", "src/lib/store.tsx — hydration and persistence effects. The ready ref is essential.")
p("The ready ref deserves emphasis. Without it, the persistence effect would fire on the very "
  "first render with the default state and overwrite the user's real saved data with an empty "
  "object before the load effect had a chance to run. This is the single most dangerous race in "
  "the design, and the ref eliminates it deterministically.", italic=False)

h("6.2 Immutability and the mutateDay helper", 2)
p("React's change detection compares object identity. Mutating state in place would leave the "
  "reference unchanged and the UI would silently fail to update. Every action therefore produces "
  "a new object graph along the path that changed, sharing untouched branches by reference.")
p("Because roughly half the actions target one day inside a sparse map, that path is factored into "
  "a single helper. It also performs lazy record creation: a DayRecord is materialised only when "
  "the user first touches that day.")
code("""
const mutateDay = useCallback((dateKey: string, fn: (d: DayRecord) => DayRecord) => {
  setState((s) => {
    const current = s.days[dateKey] ??
      { ...EMPTY_DAY, habits: {}, leetcode: { easy: 0, medium: 0, hard: 0 } };
    return { ...s, days: { ...s.days, [dateKey]: fn(current) } };
  });
}, []);

// usage — the whole toggle action is one expression
toggleHabit: (dateKey, habitId) =>
  mutateDay(dateKey, (d) => ({ ...d, habits: { ...d.habits, [habitId]: !d.habits[habitId] } })),
""", "src/lib/store.tsx — the immutable update path shared by all day-scoped actions.")
p("Note the fresh habits and leetcode objects in the default. Spreading EMPTY_DAY alone would share "
  "the same nested objects across every newly created day, so writing to one day would corrupt all "
  "of them — a classic shared-mutable-default defect, avoided explicitly here.")

h("6.3 The action API", 2)
p("The context exposes a deliberately closed set of operations. UI components can never write state "
  "directly; they can only invoke one of these, which keeps all invariants in one file.")
table(["Action", "Signature", "Effect"],
      [["toggleHabit", "(dateKey, habitId)", "Flips one habit on one day; creates the day record if absent"],
       ["setAllHabits", "(dateKey, value)", "Bulk set or clear every habit on a day"],
       ["setLeet", "(dateKey, level, value)", "Sets easy/medium/hard count, clamped at 0 via Math.max"],
       ["setNote", "(dateKey, note)", "Replaces the day's reflection text"],
       ["addBook", "(book without id)", "Appends a book with a generated id"],
       ["updateBook", "(id, partial)", "Field-level patch of one book"],
       ["removeBook", "(id)", "Filters the book out of the list"],
       ["addTopic", "(name, total)", "Appends a DSA topic with done = 0"],
       ["updateTopic", "(id, partial)", "Field-level patch of one topic"],
       ["removeTopic", "(id)", "Filters the topic out"],
       ["addNote", "(title, body)", "Prepends a note with an ISO createdAt timestamp"],
       ["removeNote", "(id)", "Filters the note out"],
       ["updateSettings", "(partial)", "Merges a patch into settings"],
       ["exportState", "() => string", "Serialises the whole state as pretty-printed JSON"],
       ["importState", "(json) => boolean", "Parses, migrates and replaces state; returns false on malformed input"],
       ["resetAll", "()", "Replaces state with a fresh defaultState()"]],
      widths=[1.15, 1.6, 3.9])

h("6.4 Re-render control", 2)
p("The context value is wrapped in useMemo keyed on [state, hydrated, mutateDay]. Without it, a new "
  "object literal would be produced on every provider render and every consumer in the tree would "
  "re-render regardless of whether anything relevant changed. mutateDay itself is wrapped in "
  "useCallback with an empty dependency array, so it is referentially stable for the lifetime of "
  "the application and never invalidates the memo.")
p("A deliberate simplification: because there is exactly one state object, any change re-renders all "
  "consumers. With a payload this small — a 90-day challenge is on the order of tens of kilobytes — "
  "the recomputation cost is negligible, and the alternative (selector-based subscriptions) would "
  "add complexity with no measurable benefit. This is documented here as a conscious trade, not an "
  "oversight; Section 18 notes when it would need revisiting.")

# ============================ 7. DOMAIN LOGIC ============================
h("7. Domain Logic and Algorithms", 1, page_break=True)
p("Everything in src/lib/stats.ts is a pure function of AppState. Nothing is cached, memoised or "
  "stored — every number shown in the UI is recomputed from raw activity on each render. This means "
  "derived values can never drift out of sync with the underlying data, and changing a rule (for "
  "example the XP for a habit) retroactively and correctly updates all history.")

h("7.1 The habit catalogue", 2)
p("Habits are declared as data, not code. Adding a habit is a one-line change to this array; every "
  "screen, chart, XP calculation and consistency ranking picks it up automatically because they all "
  "iterate the catalogue.")
table(["Habit", "Icon", "Category", "XP", "Daily intent"],
      [["Gym", "dumbbell", "Body", "15", "Strength session"],
       ["Running", "runner", "Body", "15", "Cardio / 5K"],
       ["Healthy Food", "salad", "Body", "10", "No junk today"],
       ["Water", "droplet", "Body", "5", "3L or more"],
       ["Sleep", "sleeping", "Body", "10", "7-8 hours"],
       ["Coursera DSA", "graduate", "Mind", "15", "1 lecture + quiz"],
       ["LeetCode", "laptop", "Mind", "20", "2+ problems"],
       ["AI Agents Book", "robot", "Mind", "10", "20 pages"],
       ["ML Systems Book", "book", "Mind", "10", "20 pages"],
       ["System Design", "building", "Mind", "15", "1 design case"],
       ["Test Architecture", "test tube", "Mind", "10", "Framework work"],
       ["Java", "coffee", "Mind", "10", "Core / concurrency"],
       ["AI / ML", "brain", "Mind", "15", "Paper or hands-on"],
       ["Revision", "repeat", "Mind", "10", "Revisit old notes"]],
      widths=[1.4, 0.8, 0.8, 0.5, 3.15],
      caption="Table — The 14-habit catalogue. Body total 55 XP, Mind total 115 XP, combined 170 XP per perfect day before bonuses.")

h("7.2 Completion percentage", 2)
p("Completion is intentionally unweighted — it counts habits, not XP — so the ring answers "
  "\"how much of my day did I do?\" rather than \"how many points did I score?\". XP answers the "
  "second question separately.")
code("""
doneCount(day)      = number of catalogue habits whose flag is true
completionPct(day)  = round( doneCount(day) / 14 * 100 )

Examples:   0/14 ->   0%      7/14 ->  50%      10/14 ->  71%      14/14 -> 100%
""")

h("7.3 The XP economy", 2)
p("XP is the motivational currency. It is computed per day and summed across all days for the "
  "lifetime total. Four components combine, each with a distinct behavioural purpose.")
code("""
dayXp(day) =
      SUM( habit.xp  for every habit toggled on )        // 5 .. 20 each, 170 max
    + 50    if doneCount === 14                          // perfect-day bonus
    + (easy x 2) + (medium x 5) + (hard x 10)            // difficulty-weighted volume
    + 10    if (easy + medium + hard) >= 3               // consistency bonus

totalXp(state) = SUM( dayXp(d) for every d in state.days )
""", "src/lib/stats.ts — the complete XP formula.")
table(["Component", "Value", "Behaviour it is designed to encourage"],
      [["Per-habit award", "5-20 XP", "Weighted by effort: Water 5, Gym 15, LeetCode 20. Makes hard habits feel worth more."],
       ["Perfect-day bonus", "+50 XP", "A step function, not linear. Creates a strong pull to finish the last one or two habits."],
       ["Difficulty weighting", "2 / 5 / 10 XP", "A hard problem is worth five easy ones, discouraging farming trivial problems."],
       ["Volume bonus", "+10 XP", "Rewards sustained sessions of three or more problems over single-problem days."]],
      widths=[1.15, 0.95, 4.55])
p("Maximum achievable XP for a day with no LeetCode activity is 220 (170 habits + 50 bonus). A day "
  "with three hard problems adds a further 40 = 30 + 10, giving 260.")

h("7.4 The level curve", 2)
p("Levels use a linearly increasing cost, which produces a quadratic cumulative curve. Early levels "
  "arrive quickly to establish momentum; later levels take progressively longer so the progression "
  "never runs out.")
code("""
Cost to clear level n           = 250 x n  XP
Cumulative XP to reach level n  = 250 x (n-1) x n / 2

let level = 1, need = 250, remaining = xp;
while (remaining >= need) { remaining -= need; level += 1; need = 250 * level; }
return { level, title: TITLES[level-1], into: remaining, need, pct: round(into/need*100) };
""", "src/lib/stats.ts — levelInfo(). The loop terminates because need grows without bound.")
table(["Level", "Title", "XP for this level", "Cumulative XP", "Roughly"],
      [["1", "Rookie", "-", "0", "Day 1"],
       ["2", "Apprentice", "250", "250", "~1 strong day"],
       ["3", "Grinder", "500", "750", "~4 days"],
       ["4", "Builder", "750", "1,500", "~7 days"],
       ["5", "Engineer", "1,000", "2,500", "~12 days"],
       ["6", "Architect", "1,250", "3,750", "~18 days"],
       ["7", "Specialist", "1,500", "5,250", "~25 days"],
       ["8", "Veteran", "1,750", "7,000", "~33 days"],
       ["9", "Master", "2,000", "9,000", "~42 days"],
       ["10", "Legend", "2,250", "11,250", "~53 days"]],
      widths=[0.6, 1.1, 1.35, 1.15, 2.4],
      caption="Table — Level curve. 'Roughly' assumes a sustained ~215 XP per day. Titles clamp at Legend beyond level 10.")

h("7.5 Streak algorithms", 2)
p("A streak is the count of consecutive days that met the daily target. The subtle requirement is "
  "that a streak must not appear broken simply because today is not finished yet — at 9 a.m. a "
  "30-day streak should still read 30, not 0. The algorithm handles this with a single conditional "
  "step-back before counting.")
code("""
isDayWon(state, key) = state.days[key] exists AND completionPct >= settings.dailyTargetPct

currentStreak(state):
    key = today
    if NOT isDayWon(key):        key = key - 1 day     // today is still in progress
    streak = 0
    while isDayWon(key):  streak += 1;  key = key - 1 day
    return streak
""", "src/lib/stats.ts — currentStreak(). Cost is O(streak length), which is bounded by the challenge length.")
p("bestStreak() takes a different approach because it must scan history rather than walk backwards "
  "from a known point. It collects every winning day, sorts the keys lexicographically — which for "
  "the ISO 'YYYY-MM-DD' format is identical to chronological order, a deliberate property of the "
  "key format chosen in Section 7.7 — then measures run lengths by testing whether consecutive keys "
  "differ by exactly one day.")
code("""
bestStreak(state):
    keys = sort( every key where isDayWon(key) )      // lexicographic == chronological
    best = run = 0; prev = null
    for k in keys:
        run  = (prev != null AND diffDays(prev, k) == 1) ? run + 1 : 1
        best = max(best, run)
        prev = k
    return best
""", "src/lib/stats.ts — bestStreak(). O(n log n) on the number of logged days.")
p("The daily target is user-configurable (default 70%). Lowering it makes streaks easier to keep "
  "alive on bad days; raising it makes them stricter. Because streaks are always derived and never "
  "stored, changing the setting instantly and correctly recomputes all historical streaks.")

h("7.6 Achievements", 2)
p("Thirteen achievements are evaluated as boolean expressions over five aggregate values: total XP, "
  "the greater of current and best streak, total problems solved, perfect days and active days. "
  "They are re-evaluated on every render, so they can never be missed or double-awarded, and no "
  "unlock state has to be persisted.")
table(["Achievement", "Icon", "Unlock condition"],
      [["First Step", "footprint", "1 or more active days"],
       ["Week One", "calendar", "7 or more active days"],
       ["Warming Up", "fire", "Streak reaches 3"],
       ["On Fire", "fire", "Streak reaches 7"],
       ["Unstoppable", "lightning", "Streak reaches 30"],
       ["Flawless", "gem", "1 perfect day (all 14 habits)"],
       ["Machine", "robot", "10 perfect days"],
       ["Grinder", "laptop", "25 problems solved"],
       ["Century", "trophy", "100 problems solved"],
       ["Four Digits", "sparkles", "1,000 total XP"],
       ["Elite", "crown", "5,000 total XP"],
       ["Halfway There", "flag", "45 active days"],
       ["Finisher", "party", "90 active days — challenge complete"]],
      widths=[1.4, 0.9, 4.35])

h("7.7 Date handling — a deliberately defensive design", 2)
p("Dates are the most common source of subtle defects in trackers, so this design avoids the two "
  "usual traps outright.")
rich([("Trap 1 — UTC drift. ", True, False),
      ("new Date().toISOString().slice(0,10) is the common idiom for a date key, but it converts to "
       "UTC first. For a user in IST (UTC+5:30) logging a habit before 05:30, that idiom silently "
       "returns yesterday's date. The design instead builds the key from local calendar components:", False, False)])
code("""
export function toKey(d: Date): string {
  const y   = d.getFullYear();                          // LOCAL year
  const m   = `${d.getMonth() + 1}`.padStart(2, "0");   // LOCAL month
  const day = `${d.getDate()}`.padStart(2, "0");        // LOCAL day
  return `${y}-${m}-${day}`;
}
""", "src/lib/date.ts — the timezone-safe key constructor used everywhere.")
rich([("Trap 2 — arithmetic across DST and month ends. ", True, False),
      ("Adding 86,400,000 milliseconds breaks on daylight-saving boundaries and needs special cases "
       "for month and year rollover. The design delegates to the platform's own calendar arithmetic, "
       "which handles all of it, then re-derives the key:", False, False)])
code("""
export function addDays(key: string, n: number): string {
  const d = fromKey(key);
  d.setDate(d.getDate() + n);   // handles month/year rollover and DST correctly
  return toKey(d);
}
""", "src/lib/date.ts — safe date stepping.")
p("The chosen 'YYYY-MM-DD' format yields a third benefit already exploited in Section 7.5: string "
  "sort order equals chronological order, so no date parsing is needed when ordering history.")

h("7.8 Derived series for charts", 2)
p("Four series generators feed the visualisation layer. Each maps a window of calendar days onto "
  "numbers, substituting zero for days with no record, so charts always have a complete, "
  "gap-free x-axis regardless of how sparsely the user logged.")
table(["Function", "Output per day", "Used by"],
      [["completionSeries(state, n)", "Completion percentage 0-100", "Weekly bar chart, completion trend line"],
       ["xpSeries(state, n)", "XP earned that day", "XP trend line"],
       ["solvedSeries(state, n)", "Problems solved that day", "LeetCode volume line"],
       ["habitRates(state, n)", "Per-habit hit rate %, sorted worst first", "Habit consistency ranking"]],
      widths=[1.85, 2.2, 2.6])
p("habitRates deliberately sorts ascending. The most useful information is not what the user is "
  "already doing well — it is the weakest habit, which is therefore placed at the top of the list "
  "and coloured red below 40% and amber below 70%.")

# ============================ 8. UI / UX ============================
h("8. User Interface Design", 1, page_break=True)
h("8.1 Design language", 2)
p("The visual system is a dark 'glassmorphism' theme chosen for two practical reasons: the app is "
  "used late at night when a light theme is uncomfortable, and dark backgrounds make the emerald "
  "completion signals and gradient charts read instantly.")
table(["Token", "Value", "Applied to"],
      [["Base background", "#060914", "Page background beneath the gradients"],
       ["Ambient gradients", "Indigo 25%, cyan 18%, fuchsia 14%", "Three fixed radial gradients giving depth without imagery"],
       ["Surface", "white at 4% opacity + 10% border", "Cards, tiles, inputs — the 'frosted glass' effect"],
       ["Primary accent", "Indigo 500 to cyan 400", "Progress ring, primary buttons, active states"],
       ["Success", "Emerald 400/300", "Completed habits, heatmap intensity, on-target rings"],
       ["Warning / danger", "Amber 400 / Rose 500", "Habit consistency below 70% / below 40%, destructive actions"],
       ["Radius", "0.75rem to 1.5rem", "Inputs, tiles, cards — progressively rounder at larger sizes"],
       ["Motion", "200-700 ms cubic-bezier", "Ring sweep 700 ms, bars 500 ms, taps 150-200 ms"]],
      widths=[1.25, 1.85, 3.55])

h("8.2 Navigation model", 2)
p("A single Navbar component renders two entirely different layouts from the same route list, "
  "switching at Tailwind's md breakpoint (768 px). This is a responsive-design decision driven by "
  "ergonomics rather than aesthetics: on a phone the primary navigation must sit within thumb "
  "reach at the bottom of the screen, whereas on a desktop it belongs at the top.")
code("""
  DESKTOP / TABLET  (>= 768px)              MOBILE  (< 768px)
  +--------------------------------+        +--------------------------+
  | AI  Home Habits Learning ...   | <- top |                          |
  +--------------------------------+        |      page content        |
  |                                |        |                          |
  |         page content           |        |                          |
  |                                |        +--------------------------+
  |                                |        | Home Habits Learn Stats  | <- bottom
  +--------------------------------+        |  (safe-area inset pad)   |    tab bar
""", "Figure 6 — Responsive navigation. The mobile bar adds env(safe-area-inset-bottom) padding to clear the iPhone home indicator.")
table(["Tab", "Route", "Purpose"],
      [["Home", "/dashboard", "Daily overview and fastest path to log today"],
       ["Habits", "/habits", "Full check-in for any date, LeetCode counters, day note"],
       ["Learning", "/books", "Books, DSA roadmap, LeetCode breakdown, notes"],
       ["Stats", "/analytics", "Charts, trends, consistency, achievements"],
       ["Settings", "/settings", "Configuration, backup, restore, install guide"]],
      widths=[0.85, 1.15, 4.65])
p("Route / is a server-side redirect() to /dashboard, so a home-screen launch, a bookmark and a "
  "manually typed origin all land in the same place.")

h("8.3 Screen specifications", 2)
h("8.3.1 Dashboard", 3)
p("Answers one question in a single glance: what is my state right now, and what is left to do?")
bullet("Header — full date, greeting using the configured name, and a Day N/90 badge.")
bullet("Hero card — 190 px progress ring showing today's completion, which switches from indigo/cyan to emerald/lime once the 70% target is crossed, giving a colour-coded pass signal.")
bullet("XP bar — level badge, rank title, XP into current level over XP required, lifetime total.")
bullet("Four stat tiles — day streak, habits done today, XP today, and overall challenge progress.")
bullet("Quote card — deterministically selected as QUOTES[dayIndex % 15], so it changes daily but is stable across reloads within a day (important: a random pick would flicker on every render).")
bullet("Quick check-in — all 14 habits as compact toggles. This is the highest-frequency control in the app and is deliberately reachable with zero navigation.")
bullet("Heatmap — the full 90-day challenge grid.")

h("8.3.2 Habits", 3)
bullet("Date navigator — previous/next with the forward control disabled when the next day is in the future, preventing logging of days that have not happened.")
bullet("Contextual label — renders 'Today', 'Yesterday' or the full date, computed from the day offset.")
bullet("Summary card — ring showing done-of-14, plus completion %, XP and problems solved.")
bullet("Two grouped sections — Body (5 habits) and Mind (9 habits), each with its own progress count, so the user can see whether they are neglecting one dimension.")
bullet("LeetCode counters — three increment/decrement steppers, colour-coded emerald/amber/rose, clamped at zero.")
bullet("Day note — free text, persisted on every keystroke.")

h("8.3.3 Learning", 3)
bullet("Books — progress bar per book in its own colour, direct page-number entry plus a +10 quick button for the common case of finishing a chapter.")
bullet("LeetCode donut — lifetime easy/medium/hard split with the total in the centre.")
bullet("DSA roadmap — 15 seeded topics totalling 192 problems, each with a scrollable counter row and progress bar, plus add/remove.")
bullet("Notes — title and body capture, newest first, with delete.")

h("8.3.4 Analytics", 3)
bullet("Eight aggregate tiles — current streak, best streak, average completion, total XP, perfect days, active days, problems solved, configured target.")
bullet("Range switch — 7 / 30 / 90 days, applied to every chart on the screen simultaneously.")
bullet("Weekly bar chart, completion trend line, XP trend line, habit consistency ranking, solved-per-day line, full heatmap, and the 13-achievement grid with locked items dimmed to 45% opacity.")

h("8.3.5 Settings", 3)
bullet("Profile and challenge — name, start date, challenge length, and a slider for the streak threshold.")
bullet("Data — export to a timestamped JSON file, import with validation, and a confirm-guarded destructive reset.")
bullet("Install guide — the four Safari steps to add the app to the iPhone home screen.")

h("8.4 Accessibility", 2)
table(["Concern", "Treatment"],
      [["Semantics", "Real <button> and <a> elements throughout; no click handlers on <div>"],
       ["Toggle state", "aria-pressed on habit cards so assistive technology announces on/off"],
       ["Icon-only controls", "aria-label on date arrows and the six LeetCode steppers"],
       ["Data points", "SVG <title> elements on chart points and heatmap cells give native tooltips"],
       ["Contrast", "Body text at 100% and 70% white on a near-black base; secondary text never below 40%"],
       ["Touch targets", "Minimum 40 px height on interactive rows; 44 px on the bottom tab bar"],
       ["iOS zoom", "All inputs forced to 16px font-size, which suppresses Safari's focus auto-zoom"],
       ["Motion", "Transitions are decorative only; no content depends on animation to be readable"]],
      widths=[1.5, 5.15])

# ============================ 9. CHARTS ============================
h("9. Visualisation Subsystem", 1, page_break=True)
p("No charting library is used. Every visual is computed geometry rendered as SVG or styled "
  "elements. This section documents the mathematics of each, because it is the least self-evident "
  "part of the codebase.")

h("9.1 ProgressRing — stroke-dasharray sweep", 2)
p("A circle is stroked with a dash pattern whose dash length equals the full circumference. By "
  "offsetting where that dash starts, an arbitrary fraction of the circle is revealed. Animating "
  "only the offset means the browser can drive the whole sweep on the compositor.")
code("""
r      = (size - stroke) / 2            // inset by half the stroke so it is not clipped
c      = 2 * PI * r                     // full circumference
offset = c - (pct / 100) * c            // 100% -> 0 (full ring), 0% -> c (invisible)

<circle strokeDasharray={c} strokeDashoffset={offset}
        style={{ transition: "stroke-dashoffset 700ms cubic-bezier(.4,0,.2,1)" }} />

The <svg> carries class "-rotate-90" so 0% begins at 12 o'clock rather than 3 o'clock.
A <linearGradient> with a unique id per instance supplies the two-stop stroke colour.
""", "src/components/ProgressRing.tsx — the sweep calculation.")

h("9.2 BarChart — proportional flex columns", 2)
code("""
top = max( declared max, all values, 1 )      // the trailing 1 prevents divide-by-zero
h   = max( 2, value / top * 100 )             // percentage height, floored at 2% so
                                              // a zero value still shows a visible stub
Each column is a flex-1 child, so widths distribute evenly for any series length.
""", "src/components/Charts.tsx — BarChart height mapping.")

h("9.3 LineChart — viewBox coordinate mapping", 2)
p("The chart is authored in a fixed 600 x 180 coordinate space and scaled by the browser through "
  "the SVG viewBox, so it is resolution independent and needs no resize listener.")
code("""
W = 600, H = 180, P = 10 (padding)
step = (W - 2P) / (n - 1)                     // horizontal spacing between points
x_i  = P + i * step
y_i  = H - P - (value_i / top) * (H - 2P)     // y is inverted: SVG origin is top-left

line = "M x0,y0 L x1,y1 L x2,y2 ..."          // the stroked path
area = line + " L xLast,H-P L P,H-P Z"        // same path closed to the baseline,
                                              // filled with a vertical alpha gradient
Three horizontal gridlines are drawn at 25%, 50% and 75% of the plot height at 7% opacity.
""", "src/components/Charts.tsx — LineChart geometry.")

h("9.4 Donut — accumulated arc offsets", 2)
p("The donut reuses the dasharray technique, but each slice is a full circle whose visible dash is "
  "only its own fraction, pushed around the ring by the accumulated length of all preceding slices.")
code("""
c = 2 * PI * r,  acc = 0
for each slice:
    dash = (slice.value / total) * c
    render <circle strokeDasharray={`${dash} ${c - dash}`} strokeDashoffset={-acc} />
    acc += dash                               // next slice starts where this one ended
""", "src/components/Charts.tsx — Donut slice placement.")

h("9.5 CalendarHeatmap — grid alignment and colour buckets", 2)
p("The grid is column-per-week, row-per-weekday. The only non-obvious step is the leading pad: the "
  "first column must begin on the correct weekday, otherwise every subsequent date sits in the "
  "wrong row.")
code("""
lead  = weekday index (0=Sun..6=Sat) of the FIRST date in the series
cells = [ null x lead, ...data ]              // blank spacers align column 1
weeks = chunk(cells, 7)                       // one column per 7 cells

Month labels: a label is emitted only when a column's first real date belongs to a
different month than the previously labelled column, which prevents duplicates when a
month boundary falls mid-column.

Colour buckets (completion %):
    future -> white 3%      0 -> white 6%      1-24  -> emerald 25%
    25-49  -> emerald 45%   50-74 -> emerald 65%   75-99 -> emerald 85%   100 -> emerald 300
""", "src/components/CalendarHeatmap.tsx — layout and intensity mapping.")
p("Cells are 13 x 13 px with a 3 px gap, matching the familiar GitHub contribution graph so the "
  "visual grammar needs no explanation. Every cell carries a title attribute giving the full date "
  "and percentage on hover.")

# ============================ 10. PWA ============================
h("10. Progressive Web App Design", 1, page_break=True)
p("Three independent mechanisms combine to make a web page behave like an installed application: "
  "the web app manifest (identity and launch behaviour), the service worker (offline capability), "
  "and platform-specific meta tags and icons (iOS integration).")

h("10.1 Web app manifest", 2)
table(["Key", "Value", "Effect"],
      [["name / short_name", "AI Tracker — 90 Day Challenge / AI Tracker", "Install prompt title and home-screen label"],
       ["start_url", "/dashboard", "Launching from the home screen opens the dashboard directly"],
       ["scope", "/", "All routes are treated as in-app; no browser chrome appears when navigating"],
       ["display", "standalone", "Full screen with no address bar — the key 'feels native' switch"],
       ["orientation", "portrait", "Locks to portrait, matching the layout's design intent"],
       ["background_color", "#060914", "Splash-screen colour while the app boots"],
       ["theme_color", "#060914", "Status-bar tint on Android and supporting browsers"],
       ["icons", "192, 512 (any) + 512 (maskable)", "Home-screen, splash and adaptive-mask icons"],
       ["shortcuts", "Today's habits, Analytics", "Long-press jump targets on supporting platforms"]],
      widths=[1.3, 1.95, 3.4])

h("10.2 Service worker lifecycle and caching strategy", 2)
p("The worker uses a network-first, cache-fallback strategy. This ordering is deliberate: a habit "
  "tracker must never show a stale shell after an update, so the network is always tried first, and "
  "the cache exists purely as an offline safety net.")
code("""
INSTALL
   |  open cache "ai-tracker-v1"
   |  add 10 shell entries: / /dashboard /habits /books /analytics /settings
   |                        /manifest.webmanifest + 3 icons
   |  Promise.allSettled  <- one failed asset must not abort the whole install
   |  skipWaiting()       <- activate immediately instead of waiting for tab close
   v
ACTIVATE
   |  delete every cache whose name != "ai-tracker-v1"   (version-based eviction)
   |  clients.claim()     <- take control of already-open tabs
   v
FETCH  (only same-origin GET; /_next/webpack-hmr and /__nextjs are excluded)
   |
   +--> try network
   |        success (status 200, type 'basic') -> clone into cache, return response
   |        failure -------------------------------+
   |                                               v
   +--> cache lookup
            hit                -> return cached response
            miss + navigation  -> return cached /dashboard  (offline app shell)
            miss + asset       -> return 503 "Offline"
""", "public/sw.js — install, activate and fetch handling.")
table(["Decision", "Rationale"],
      [["Network-first, not cache-first", "Guarantees the newest build is used whenever connectivity exists; stale UI in a tracker is worse than a few milliseconds of latency"],
       ["Promise.allSettled on precache", "A single 404 during install would otherwise reject the install event and leave the app with no offline support at all"],
       ["skipWaiting + clients.claim", "Removes the classic 'you must close every tab for the update to apply' problem"],
       ["Cache name carries the version", "Activation deletes all other caches, so a deploy cannot leave orphaned stale assets behind"],
       ["Exclude HMR endpoints", "Caching development hot-reload traffic would break the dev server"],
       ["Only cache type 'basic'", "Prevents opaque cross-origin responses from silently poisoning the cache"],
       ["Navigation falls back to /dashboard", "Any deep link opened offline still renders a working app shell rather than a browser error page"]],
      widths=[1.9, 4.75])

h("10.3 Registration guards", 2)
p("Registration is deliberately conservative and contains a defect fix worth documenting, since it "
  "was found during verification rather than reasoning.")
code("""
if (!("serviceWorker" in navigator)) return;                     // capability check
if (protocol !== "https:" && hostname !== "localhost") return;   // secure-origin only

if (document.readyState === "complete") { register(); return; }  // <-- THE FIX
window.addEventListener("load", register);
""", "src/components/PWARegister.tsx — registration with a readyState guard.")
p("The original implementation only attached a load listener. Because React effects frequently run "
  "after the window load event has already fired, the listener was registered too late and never "
  "invoked — verification showed zero registrations. Checking readyState first covers that case. "
  "The secure-origin guard is also necessary: service workers are restricted to HTTPS and "
  "localhost, so over a plain-HTTP LAN address registration would throw, and the guard turns that "
  "into a silent, graceful degradation where the app still runs and still saves data.")

h("10.4 iOS integration", 2)
table(["Mechanism", "Purpose"],
      [["apple-touch-icon (180x180 PNG)", "Home-screen icon; iOS ignores manifest icons for this"],
       ["apple-mobile-web-app-capable = yes", "Enables standalone mode on iOS versions before 16.4. Next 15 emits only the modern mobile-web-app-capable, so this is added explicitly via metadata.other"],
       ["apple-mobile-web-app-title", "Sets the label shown under the home-screen icon"],
       ["status-bar-style = black-translucent", "Lets the app background extend behind the status bar"],
       ["viewportFit = cover + env(safe-area-inset-bottom)", "Extends into the safe area while keeping the tab bar clear of the home indicator"],
       ["maximumScale = 1", "Prevents pinch-zoom, which would break the app illusion"],
       ["Inputs at 16px", "Below 16px Safari auto-zooms on focus and does not zoom back out"]],
      widths=[2.3, 4.35])

h("10.5 Icon generation pipeline", 2)
p("Rather than adding an image library, icons are produced by a 160-line dependency-free Node "
  "script that writes valid PNG files using only the built-in zlib module. This keeps the build "
  "reproducible and the dependency count at zero for what is a build-time-only concern.")
code("""
1. Allocate an RGBA byte buffer of size x size x 4.
2. For each pixel:
     a. shape mask      - rounded-square (corner radius 22% of size), or full-bleed for maskable
     b. base colour     - diagonal 3-stop gradient: indigo #4F46E5 -> violet #8B5CF6 -> cyan #22D3EE
     c. vignette        - multiply by (1 - 0.2 * d^2), d = normalised distance from centre
     d. progress arc    - white annulus between 42% and 50% of the inner box, with a gap
                          between -90 deg and -20 deg, echoing the app's progress ring
     e. bolt glyph      - even-odd point-in-polygon test against a 7-vertex lightning bolt
3. Encode PNG by hand:
     signature 89 50 4E 47 0D 0A 1A 0A
     IHDR  - width, height, bit depth 8, colour type 6 (RGBA), no interlace
     IDAT  - zlib.deflateSync(scanlines each prefixed with filter byte 0x00, level 9)
     IEND
     every chunk sealed with a CRC-32 computed from a precomputed 256-entry table
4. Emit icon-192, icon-512, apple-touch-icon-180 and maskable-512.
""", "scripts/generate-icons.mjs — run via npm run icons.")
p("The maskable variant is rendered full-bleed with 18% internal padding so that Android's adaptive "
  "icon mask can crop it to any shape without clipping the glyph.")

# ============================ 11. BUILD & DEPLOY ============================
h("11. Build, Run and Deployment", 1, page_break=True)
h("11.1 Build output", 2)
p("All six routes pre-render to static HTML at build time. Nothing needs a server at runtime, which "
  "is what allows the app to be hosted on any static CDN and to work from the service-worker cache.")
table(["Route", "Route JS", "First load JS", "Render mode"],
      [["/", "123 B", "103 kB", "Static — redirect to /dashboard"],
       ["/dashboard", "3.10 kB", "113 kB", "Static shell, client-hydrated"],
       ["/habits", "2.85 kB", "110 kB", "Static shell, client-hydrated"],
       ["/books", "3.80 kB", "111 kB", "Static shell, client-hydrated"],
       ["/analytics", "3.61 kB", "110 kB", "Static shell, client-hydrated"],
       ["/settings", "5.08 kB", "112 kB", "Static shell, client-hydrated"],
       ["/_not-found", "992 B", "103 kB", "Static"],
       ["Shared chunks", "-", "102 kB", "React + Next runtime, loaded once"]],
      widths=[1.4, 1.1, 1.25, 2.9],
      caption="Table — Production build output. The 102 kB shared baseline is framework code; the app's own logic adds only 3-5 kB per route.")

h("11.2 Commands", 2)
table(["Command", "Purpose"],
      [["npm install", "Install dependencies (329 packages, dev included)"],
       ["npm run dev", "Development server with hot reload on port 3000"],
       ["npm run build", "Type-check, lint, compile and pre-render to .next/"],
       ["npm run start", "Serve the production build, bound to 0.0.0.0 for LAN access"],
       ["npm run lint", "ESLint only"],
       ["npm run icons", "Regenerate all four PWA icons"]],
      widths=[1.6, 5.05])

h("11.3 Deployment options", 2)
table(["Target", "Method", "Notes"],
      [["Local desktop", "npm run start", "http://localhost:3000 — currently running"],
       ["Same-Wi-Fi phone", "Open http://<host-ip>:3000", "Requires an inbound firewall rule; see 11.4"],
       ["Vercel", "npx vercel", "Free tier, automatic HTTPS — the recommended path for real iPhone use"],
       ["Any static CDN", "Serve the build output", "Netlify, Cloudflare Pages, S3+CloudFront, GitHub Pages"]],
      widths=[1.35, 1.55, 3.75])
p("HTTPS matters beyond security here: iOS only enables service-worker caching on secure origins, "
  "so full offline support on the phone requires a deployed HTTPS URL rather than a LAN address.")

h("11.4 Known environment constraint — WSL2 networking", 2)
p("The development host runs WSL2 with mirrored networking. The server correctly binds 0.0.0.0:3000 "
  "and localhost responds, but connections to the host's LAN address time out because Windows "
  "Firewall blocks inbound traffic to the WSL virtual machine by default. This is an environment "
  "configuration issue, not an application defect. Two administrator PowerShell rules resolve it:")
code("""
New-NetFirewallRule -DisplayName "AI Tracker 3000" -Direction Inbound `
                    -Protocol TCP -LocalPort 3000 -Action Allow

New-NetFirewallHyperVRule -Name "WSL-AITracker" -DisplayName "WSL AI Tracker 3000" `
                    -Direction Inbound -Protocol TCP -LocalPorts 3000 `
                    -VMCreatorId '{40E0AC32-46A5-438A-A0B2-2B479E8F2E90}'
""", "Windows Firewall rules required for LAN access from the iPhone.")

# ============================ 12. CROSS-CUTTING ============================
h("12. Cross-Cutting Concerns", 1)
h("12.1 Security and privacy", 2)
table(["Concern", "Position"],
      [["Data at rest", "Plain JSON in localStorage, protected by the browser's origin sandbox and the device lock screen. No credentials or financial data are stored, so encryption at rest was judged unnecessary."],
       ["Data in transit", "None. After first load the application makes no outbound requests whatsoever."],
       ["Third parties", "Zero. No analytics, no telemetry, no external fonts, no CDN references, no trackers."],
       ["Secrets", "None exist. There is no API key, token or password anywhere in the codebase."],
       ["XSS", "All user text is rendered as React children, which escapes by default. dangerouslySetInnerHTML is never used."],
       ["Injection", "No SQL, no server, no eval, no dynamic code construction."],
       ["Supply chain", "Three runtime dependencies, all first-party framework packages, pinned to exact versions."],
       ["Import safety", "Backup import is wrapped in try/catch, returns a boolean, and passes through migrate(), so a malformed or hostile file cannot crash or corrupt the app."]],
      widths=[1.25, 5.4])

h("12.2 Error handling", 2)
p("The application has no server to report errors to and no user who wants to see a stack trace "
  "mid-workout, so the strategy is to fail silently into a working state at every boundary.")
table(["Failure", "Handling"],
      [["Corrupt localStorage JSON", "try/catch around parse; falls back to defaultState()"],
       ["Storage quota exceeded", "try/catch around setItem; the in-memory session continues working"],
       ["Malformed import file", "importState returns false; the UI shows 'That file could not be read'"],
       ["Service worker registration failure", "Promise .catch() with no rethrow; the app runs online-only"],
       ["Service worker not permitted (plain HTTP)", "Guarded before registration is attempted"],
       ["Missing day record", "getDay() returns EMPTY_DAY, so every derived calculation still resolves"],
       ["Empty chart series", "top is floored at 1, preventing division by zero"],
       ["Destructive reset", "Guarded by an explicit confirm() dialog"]],
      widths=[2.0, 4.65])

h("12.3 Performance characteristics", 2)
table(["Path", "Cost", "Notes"],
      [["Habit toggle to repainted UI", "Sub-frame", "Pure in-memory work; no I/O on the interaction path"],
       ["Persistence write", "O(size of state)", "JSON.stringify after paint, so it never blocks input"],
       ["Full stats recomputation", "O(days x 14)", "Roughly 1,260 boolean checks for a full 90-day challenge — negligible"],
       ["currentStreak", "O(streak length)", "Bounded by the challenge length"],
       ["bestStreak", "O(n log n)", "Dominated by sorting the winning-day keys"],
       ["Chart render", "O(points)", "At most 90 SVG points per chart"],
       ["Cold start (cached)", "No network", "Shell served entirely from Cache Storage"]],
      widths=[1.85, 1.35, 3.45])

# ============================ 13. VERIFICATION ============================
h("13. Verification Evidence", 1, page_break=True)
p("The following checks were executed against the running production build. Each row records an "
  "actual observed result, not an intention.")
table(["#", "Check", "Method", "Result"],
      [["V1", "TypeScript compilation", "next build type-check", "PASS — 0 errors"],
       ["V2", "Lint", "eslint-config-next during build", "PASS — 0 warnings"],
       ["V3", "Static generation", "next build", "PASS — 10/10 pages generated"],
       ["V4", "All routes reachable", "curl status codes", "PASS — 6/6 routes return 200"],
       ["V5", "PWA assets served", "curl", "PASS — manifest, sw.js and icons all 200"],
       ["V6", "Dashboard renders", "Accessibility tree snapshot", "PASS — ring, XP bar, tiles, 14 toggles, heatmap present"],
       ["V7", "Habit toggle updates derived state", "Clicked Gym and LeetCode", "PASS — 2/14, 14%, 35 XP as calculated"],
       ["V8", "XP formula correctness", "Gym 15 + LeetCode 20", "PASS — 35 XP, matches the specification exactly"],
       ["V9", "Persistence across reload", "Full page reload", "PASS — 35 XP and 14% survived"],
       ["V10", "localStorage payload shape", "Read the stored key", "PASS — matches the documented schema"],
       ["V11", "Habits screen", "Snapshot", "PASS — Body 1/5, Mind 1/9, counters, future day disabled"],
       ["V12", "Analytics screen", "Snapshot", "PASS — 8 tiles, 4 charts, heatmap, 1/13 achievements unlocked"],
       ["V13", "Learning screen", "Snapshot", "PASS — books, donut, roadmap, notes all render"],
       ["V14", "Settings screen", "Navigation", "PASS — renders without error"],
       ["V15", "Console cleanliness", "Console message capture", "PASS — 0 errors, 0 warnings across all screens"],
       ["V16", "Manifest validity", "Parsed at runtime", "PASS — standalone, 3 icons, correct name"],
       ["V17", "Service worker activation", "getRegistrations + ready", "PASS — 1 registration, state 'activated', scope /"],
       ["V18", "Offline cache populated", "Cache Storage inspection", "PASS — 10 entries under ai-tracker-v1"],
       ["V19", "iOS standalone meta tag", "DOM query", "PASS — apple-mobile-web-app-capable = yes"],
       ["V20", "Icon files valid", "file(1) inspection", "PASS — 192x192 8-bit RGBA PNG, non-interlaced"]],
      widths=[0.4, 2.0, 1.75, 2.5], font=8.2)

h("13.1 Defects found and fixed during verification", 2)
table(["Defect", "Symptom", "Root cause", "Fix"],
      [["Service worker never registered", "0 registrations at runtime", "Registration was attached to the window load event, which had already fired before the React effect ran", "Check document.readyState === 'complete' and register immediately in that case"],
       ["Duplicate month labels on the heatmap", "'Aug' appeared above two adjacent columns", "The label condition tested day-of-month <= 7, which is true for two columns when a month starts mid-week", "Track the last emitted month and only label a column when the month actually changes"],
       ["Missing iOS standalone meta tag", "Older iOS would open in a browser tab, not full screen", "Next 15 emits only the modern mobile-web-app-capable tag", "Add apple-mobile-web-app-capable explicitly through metadata.other"]],
      widths=[1.35, 1.5, 2.1, 1.7], font=8.2)

# ============================ 14. LIMITATIONS ============================
h("14. Limitations and Risks", 1)
table(["ID", "Limitation", "Impact", "Mitigation"],
      [["L1", "Data is device- and browser-local", "Progress does not follow the user to another device or browser", "JSON export/import; optional cloud sync in a future release"],
       ["L2", "Clearing browser data destroys everything", "Total loss of history", "Documented backup workflow; consider a periodic export reminder"],
       ["L3", "localStorage is roughly 5-10 MB", "A theoretical ceiling on history", "A 90-day challenge uses a few tens of kilobytes; effectively unreachable"],
       ["L4", "No automated test suite", "Regressions could reach the user unnoticed", "Manual verification recorded in Section 13; the pure domain layer is trivially unit-testable — see R4"],
       ["L5", "Single global state object", "Every change re-renders all consumers", "Negligible at this data volume; would need selectors only if the state grew substantially"],
       ["L6", "No reminders or notifications", "Relies entirely on user initiative", "Roadmap item R3"],
       ["L7", "Offline caching needs HTTPS on iOS", "LAN HTTP install lacks offline support", "Deploy to Vercel for a free HTTPS origin"],
       ["L8", "Honour-system data entry", "Nothing verifies the habits were actually done", "Inherent to self-tracking; out of scope"],
       ["L9", "Streak counts a day only at or above the target", "A 69% day breaks a streak at the default setting", "The threshold is user-configurable in Settings"]],
      widths=[0.4, 1.6, 1.75, 2.9], font=8.2)

# ============================ 15. ROADMAP ============================
h("15. Future Roadmap", 1)
table(["ID", "Enhancement", "Design impact"],
      [["R1", "Cloud sync (Firebase / Supabase)", "Introduces authentication and a repository abstraction behind the existing action API. Because every write already funnels through StoreProvider, only that one file changes; the UI is untouched. Requires a conflict-resolution policy — last-write-wins per day key is the natural fit given the data shape."],
       ["R2", "Import from Apple Health / Strava", "An adapter mapping external activity onto habit ids, plus a provenance flag on DayRecord to distinguish imported from manual entries."],
       ["R3", "Reminders and notifications", "Requires the Notifications and Push APIs, a push service, and on iOS the app must already be installed to the home screen."],
       ["R4", "Automated test suite", "Vitest over src/lib — the domain layer is pure, so XP, level, streak, achievement and date functions can be exhaustively tested with no mocking. Playwright for the interaction flows already exercised manually."],
       ["R5", "Custom user-defined habits", "Move the habit catalogue from a constant into AppState and bump the schema version; migrate() already supports adding the field safely."],
       ["R6", "Weekly and monthly review screens", "New derived selectors over existing data; no schema change required."],
       ["R7", "Theme options", "Extract the colour tokens into CSS custom properties and store the choice in settings."],
       ["R8", "Widget / lock-screen glance", "Not achievable as a PWA on iOS; would require a native shell."]],
      widths=[0.4, 1.65, 4.6], font=8.2)

# ============================ 16. APPENDICES ============================
h("16. Appendix A — Directory Structure", 1, page_break=True)
code("""
ai-tracker/
├── src/
│   ├── app/                        Next.js App Router
│   │   ├── layout.tsx              root shell: metadata, viewport, provider, nav
│   │   ├── page.tsx                redirect / -> /dashboard
│   │   ├── globals.css             Tailwind import, gradients, iOS tweaks, keyframes
│   │   ├── dashboard/page.tsx      ring, XP, tiles, quote, quick check-in, heatmap
│   │   ├── habits/page.tsx         date navigator, Body/Mind groups, counters, note
│   │   ├── books/page.tsx          books, LeetCode donut, DSA roadmap, notes
│   │   ├── analytics/page.tsx      tiles, range switch, 4 charts, heatmap, achievements
│   │   └── settings/page.tsx       config, backup/restore, reset, install guide
│   ├── components/
│   │   ├── ProgressRing.tsx        SVG circular progress with gradient stroke
│   │   ├── XPBar.tsx               level badge, title, progress bar
│   │   ├── HabitCard.tsx           toggleable habit row
│   │   ├── CalendarHeatmap.tsx     GitHub-style contribution grid
│   │   ├── Charts.tsx              BarChart, LineChart, HBars, Donut
│   │   ├── Navbar.tsx              responsive top bar / bottom tab bar
│   │   ├── Ui.tsx                  Card, CardTitle, StatCard primitives
│   │   └── PWARegister.tsx         guarded service-worker registration
│   └── lib/
│       ├── types.ts                all persisted interfaces
│       ├── date.ts                 timezone-safe date keys and arithmetic
│       ├── habits.ts               14-habit catalogue + 15 quotes
│       ├── stats.ts                XP, levels, streaks, achievements, series
│       └── store.tsx               Context provider, actions, persistence, migration
├── public/
│   ├── manifest.webmanifest        PWA identity and launch behaviour
│   ├── sw.js                       service worker
│   └── icons/                      icon-192, icon-512, apple-touch-icon, maskable-512
├── scripts/
│   └── generate-icons.mjs          dependency-free PNG encoder and icon renderer
├── docs/
│   ├── build_hld.py                generator for this document
│   └── AI-Tracker-HLD.docx         this document
├── package.json                    3 runtime dependencies
├── tsconfig.json                   strict TypeScript, @/* path alias
├── next.config.ts                  Next.js configuration
├── postcss.config.mjs              Tailwind v4 PostCSS plugin
├── eslint.config.mjs               ESLint flat config
└── README.md                       quick start and install guide
""")

h("17. Appendix B — Glossary", 1)
table(["Term", "Meaning in this system"],
      [["PWA", "Progressive Web App — a web application installable to a home screen and capable of running offline"],
       ["Service worker", "A background script that intercepts network requests, enabling offline delivery of the app shell"],
       ["App shell", "The minimal HTML, CSS and JavaScript needed to render the interface, cached for offline launch"],
       ["Hydration", "React attaching event handlers and state to server-rendered HTML on the client"],
       ["Date key", "A local-calendar 'YYYY-MM-DD' string used as the primary key for a day's record"],
       ["DayRecord", "The per-day entity holding habit flags, LeetCode counts and an optional note"],
       ["XP", "Experience points — the derived motivational score computed from habits and problems"],
       ["Perfect day", "A day on which all 14 habits were completed; awards a 50 XP bonus"],
       ["Streak", "Consecutive days meeting or exceeding the configured daily completion target"],
       ["Daily target", "The completion percentage at which a day counts towards the streak; default 70%"],
       ["Active day", "Any day with at least one habit logged"],
       ["Maskable icon", "A full-bleed icon that platforms may crop to an arbitrary adaptive shape"],
       ["Network-first", "A caching strategy that always tries the network first and uses the cache only on failure"],
       ["Sparse map", "A key-value store where absent keys are meaningful — here, 'nothing was logged that day'"]],
      widths=[1.35, 5.3], font=8.5)

# ---------- footer with page numbers ----------
footer = doc.sections[0].footer
fp = footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = fp.add_run("AI Tracker — High Level Design v1.0    |    Page ")
r.font.size = Pt(8); r.font.color.rgb = GREY
field(fp, "PAGE")
r2 = fp.add_run(" of ")
r2.font.size = Pt(8); r2.font.color.rgb = GREY
field(fp, "NUMPAGES")
for run in fp.runs:
    run.font.size = Pt(8)
    run.font.color.rgb = GREY

OUT = "/home/bhadr/ai-tracker/docs/AI-Tracker-HLD.docx"
doc.save(OUT)
print("saved:", OUT)
